from __future__ import annotations

import json
import re
import uuid
from collections import Counter
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from pathlib import Path

import numpy as np
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

MEMORY_SOURCE_PREFIX = "memory://"
SUPPORTED_DOCUMENT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".rst",
    ".pdf",
    ".docx",
}
_STOPWORDS = {
    "a",
    "as",
    "o",
    "os",
    "de",
    "da",
    "do",
    "das",
    "dos",
    "e",
    "em",
    "no",
    "na",
    "nos",
    "nas",
    "por",
    "para",
    "com",
    "sem",
    "um",
    "uma",
    "uns",
    "umas",
    "que",
    "se",
    "ao",
    "aos",
    "à",
    "às",
    "the",
    "and",
    "or",
    "of",
    "to",
    "for",
    "in",
    "on",
    "at",
    "is",
    "are",
    "was",
    "were",
}


@dataclass
class RAGChunk:
    chunk_id: str
    source: str
    content: str
    embedding: list[float]
    metadata: dict[str, str] | None = None


@dataclass
class RAGSearchResult:
    source: str
    content: str
    score: float
    metadata: dict[str, str] | None = None


@dataclass
class MemoryRecord:
    source: str
    content: str
    metadata: dict[str, str]


@dataclass
class DocumentDebugReport:
    path: str
    extension: str
    file_size: int
    extracted_chars: int
    chunk_count: int
    printable_ratio: float
    replacement_chars: int
    warnings: list[str]
    sample: str


@dataclass
class RAGDebugSearchResult:
    chunk_id: str
    source: str
    score: float
    content_preview: str
    metadata: dict[str, str] | None = None


class OpenAIRAGStore:
    """Simple JSON-backed vector store for OpenAI embedding based RAG."""

    def __init__(
        self,
        index_path: Path,
        embeddings: OpenAIEmbeddings,
        chunk_size: int = 700,
        chunk_overlap: int = 180,
    ) -> None:
        self.index_path = index_path
        self._embeddings = embeddings
        self._chunks: list[RAGChunk] = []
        self._chunk_size = chunk_size
        self._chunk_overlap = chunk_overlap
        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", " ", ""],
        )

    def load(self) -> None:
        if not self.index_path.exists():
            self._chunks = []
            return

        raw = json.loads(self.index_path.read_text(encoding="utf-8"))
        chunks: list[RAGChunk] = []
        for item in raw.get("chunks", []):
            source = str(item.get("source", "")).strip()
            content = str(item.get("content", "")).strip()
            embedding_raw = item.get("embedding", [])
            if not source or not content or not isinstance(embedding_raw, list):
                continue
            try:
                embedding = [float(value) for value in embedding_raw]
            except (TypeError, ValueError):
                continue
            metadata = _coerce_metadata(item.get("metadata"))
            chunk_id = str(item.get("chunk_id", uuid.uuid4().hex))
            chunks.append(
                RAGChunk(
                    chunk_id=chunk_id,
                    source=source,
                    content=content,
                    embedding=embedding,
                    metadata=metadata,
                )
            )
        self._chunks = chunks

    def save(self) -> None:
        self.index_path.parent.mkdir(parents=True, exist_ok=True)
        payload = {"version": 2, "chunks": [asdict(chunk) for chunk in self._chunks]}
        self.index_path.write_text(json.dumps(payload), encoding="utf-8")

    def clear(self) -> None:
        self._chunks = []
        self.save()

    def clear_documents(self) -> None:
        self._chunks = [chunk for chunk in self._chunks if _is_memory_source(chunk.source)]
        self.save()

    def clear_memories(self) -> None:
        self._chunks = [chunk for chunk in self._chunks if not _is_memory_source(chunk.source)]
        self.save()

    def count(self) -> int:
        return len(self._chunks)

    def count_documents(self) -> int:
        return len({chunk.source for chunk in self._chunks if not _is_memory_source(chunk.source)})

    def count_memories(self) -> int:
        return len({chunk.source for chunk in self._chunks if _is_memory_source(chunk.source)})

    def list_sources(self, include_memories: bool = False) -> list[str]:
        if include_memories:
            return sorted({chunk.source for chunk in self._chunks})
        return sorted({chunk.source for chunk in self._chunks if not _is_memory_source(chunk.source)})

    def list_memories(self, limit: int = 30) -> list[MemoryRecord]:
        grouped_contents: dict[str, list[str]] = {}
        grouped_metadata: dict[str, dict[str, str]] = {}
        for chunk in self._chunks:
            if not _is_memory_source(chunk.source):
                continue
            grouped_contents.setdefault(chunk.source, []).append(chunk.content.strip())
            if chunk.source not in grouped_metadata and chunk.metadata:
                grouped_metadata[chunk.source] = dict(chunk.metadata)

        records: list[MemoryRecord] = []
        for source, parts in grouped_contents.items():
            metadata = grouped_metadata.get(source, {"kind": "memory"})
            content = " ".join(part for part in parts if part).strip()
            records.append(MemoryRecord(source=source, content=content, metadata=metadata))

        records.sort(key=lambda item: item.metadata.get("created_at", ""), reverse=True)
        if limit <= 0:
            return []
        return records[:limit]

    def add_document(self, document_path: Path) -> int:
        source = str(document_path.resolve())
        text, warnings = _read_document_text(document_path, strict_quality=True)
        metadata: dict[str, str] = {
            "kind": "document",
            "extension": document_path.suffix.lower(),
        }
        if warnings:
            metadata["warnings"] = " | ".join(warnings[:3])
        return self.add_text(
            source=source,
            text=text,
            metadata=metadata,
            replace_source=True,
            chunk_id_prefix=document_path.stem,
        )

    def inspect_document(self, document_path: Path) -> DocumentDebugReport:
        return self.inspect_document_path(
            document_path,
            chunk_size=self._chunk_size,
            chunk_overlap=self._chunk_overlap,
        )

    @staticmethod
    def inspect_document_path(
        document_path: Path,
        chunk_size: int = 900,
        chunk_overlap: int = 120,
    ) -> DocumentDebugReport:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", " ", ""],
        )
        text, warnings = _read_document_text(document_path, strict_quality=False)
        normalized = text.strip()
        chunks = splitter.split_text(normalized) if normalized else []
        printable_ratio, replacement_chars = _text_quality(normalized)
        sample = normalized.replace("\n", " ")
        if len(sample) > 320:
            sample = f"{sample[:320]}..."
        return DocumentDebugReport(
            path=str(document_path.resolve()),
            extension=document_path.suffix.lower(),
            file_size=document_path.stat().st_size,
            extracted_chars=len(normalized),
            chunk_count=len(chunks),
            printable_ratio=printable_ratio,
            replacement_chars=replacement_chars,
            warnings=warnings,
            sample=sample,
        )

    def add_memory(self, note: str, topic: str | None = None, importance: str = "medium") -> str:
        normalized_note = note.strip()
        if not normalized_note:
            raise ValueError("Memory note cannot be empty.")

        topic_label = (topic or "general").strip() or "general"
        topic_slug = _slugify(topic_label)
        source = f"{MEMORY_SOURCE_PREFIX}{topic_slug}/{uuid.uuid4().hex[:12]}"
        metadata = {
            "kind": "memory",
            "topic": topic_label,
            "importance": _normalize_importance(importance),
            "created_at": datetime.now(timezone.utc).isoformat(),
        }
        self.add_text(
            source=source,
            text=normalized_note,
            metadata=metadata,
            replace_source=False,
            chunk_id_prefix=f"memory-{topic_slug}",
        )
        return source

    def add_text(
        self,
        source: str,
        text: str,
        metadata: dict[str, str] | None = None,
        replace_source: bool = False,
        chunk_id_prefix: str = "chunk",
    ) -> int:
        normalized_text = text.strip()
        if not normalized_text:
            return 0

        chunks = self._splitter.split_text(normalized_text)
        if not chunks:
            return 0

        vectors = self._embeddings.embed_documents(chunks)
        if replace_source:
            self._chunks = [chunk for chunk in self._chunks if chunk.source != source]

        prefix = _slugify(chunk_id_prefix)
        for index, (content, embedding) in enumerate(zip(chunks, vectors)):
            self._chunks.append(
                RAGChunk(
                    chunk_id=f"{prefix}-{index}-{uuid.uuid4().hex[:8]}",
                    source=source,
                    content=content,
                    embedding=embedding,
                    metadata=dict(metadata) if metadata else None,
                )
            )
        self.save()
        return len(chunks)

    def search(
        self,
        query: str,
        k: int = 4,
        source_prefix: str | None = None,
        exclude_source_prefix: str | None = None,
        min_score: float | None = None,
    ) -> list[RAGSearchResult]:
        query = query.strip()
        if not query:
            return []

        filtered_chunks = [
            chunk
            for chunk in self._chunks
            if _match_source_filters(
                chunk.source,
                source_prefix=source_prefix,
                exclude_source_prefix=exclude_source_prefix,
            )
        ]
        if not filtered_chunks:
            return []

        query_vector = np.array(self._embeddings.embed_query(query), dtype=np.float32)
        matrix = np.array([chunk.embedding for chunk in filtered_chunks], dtype=np.float32)

        matrix_norm = np.linalg.norm(matrix, axis=1) + 1e-12
        query_norm = np.linalg.norm(query_vector) + 1e-12
        similarities = (matrix @ query_vector) / (matrix_norm * query_norm)

        top_indices = np.argsort(similarities)[::-1][: max(k, 1)]
        results: list[RAGSearchResult] = []
        for idx in top_indices:
            vector_score = float(similarities[int(idx)])
            if min_score is not None and vector_score < min_score:
                continue
            chunk = filtered_chunks[int(idx)]
            metadata = dict(chunk.metadata) if chunk.metadata else {}
            metadata["chunk_id"] = chunk.chunk_id
            metadata["vector_score"] = f"{vector_score:.6f}"
            metadata["lexical_score"] = "0.000000"
            metadata["final_score"] = f"{vector_score:.6f}"
            results.append(
                RAGSearchResult(
                    source=chunk.source,
                    content=chunk.content,
                    score=vector_score,
                    metadata=metadata,
                )
            )
        return results

    def search_hybrid(
        self,
        query: str,
        k: int = 8,
        source_prefix: str | None = None,
        exclude_source_prefix: str | None = None,
        min_score: float | None = None,
    ) -> list[RAGSearchResult]:
        query = query.strip()
        if not query:
            return []

        filtered_chunks = [
            chunk
            for chunk in self._chunks
            if _match_source_filters(
                chunk.source,
                source_prefix=source_prefix,
                exclude_source_prefix=exclude_source_prefix,
            )
        ]
        if not filtered_chunks:
            return []

        query_vector = np.array(self._embeddings.embed_query(query), dtype=np.float32)
        matrix = np.array([chunk.embedding for chunk in filtered_chunks], dtype=np.float32)
        matrix_norm = np.linalg.norm(matrix, axis=1) + 1e-12
        query_norm = np.linalg.norm(query_vector) + 1e-12
        vector_scores = (matrix @ query_vector) / (matrix_norm * query_norm)

        lexical_scores = _compute_lexical_scores(query, [chunk.content for chunk in filtered_chunks])
        max_lexical = max(lexical_scores) if lexical_scores else 0.0

        ranked: list[tuple[float, float, float, int]] = []
        for idx, vector_score in enumerate(vector_scores):
            lexical_raw = lexical_scores[idx] if idx < len(lexical_scores) else 0.0
            lexical_norm = lexical_raw / max(max_lexical, 1e-9)
            final_score = (0.65 * float(vector_score)) + (0.35 * lexical_norm)
            ranked.append((final_score, float(vector_score), lexical_norm, idx))

        ranked.sort(key=lambda item: item[0], reverse=True)
        results: list[RAGSearchResult] = []
        for final_score, vector_score, lexical_norm, idx in ranked[: max(k, 1)]:
            if min_score is not None and final_score < min_score:
                continue
            chunk = filtered_chunks[idx]
            metadata = dict(chunk.metadata) if chunk.metadata else {}
            metadata["chunk_id"] = chunk.chunk_id
            metadata["vector_score"] = f"{vector_score:.6f}"
            metadata["lexical_score"] = f"{lexical_norm:.6f}"
            metadata["final_score"] = f"{final_score:.6f}"
            results.append(
                RAGSearchResult(
                    source=chunk.source,
                    content=chunk.content,
                    score=final_score,
                    metadata=metadata,
                )
            )
        return results

    def debug_search(
        self,
        query: str,
        k: int = 6,
        source_prefix: str | None = None,
        exclude_source_prefix: str | None = None,
        min_score: float | None = None,
    ) -> list[RAGDebugSearchResult]:
        query = query.strip()
        if not query:
            return []

        filtered_chunks = [
            chunk
            for chunk in self._chunks
            if _match_source_filters(
                chunk.source,
                source_prefix=source_prefix,
                exclude_source_prefix=exclude_source_prefix,
            )
        ]
        if not filtered_chunks:
            return []

        query_vector = np.array(self._embeddings.embed_query(query), dtype=np.float32)
        matrix = np.array([chunk.embedding for chunk in filtered_chunks], dtype=np.float32)
        matrix_norm = np.linalg.norm(matrix, axis=1) + 1e-12
        query_norm = np.linalg.norm(query_vector) + 1e-12
        vector_scores = (matrix @ query_vector) / (matrix_norm * query_norm)

        lexical_scores = _compute_lexical_scores(query, [chunk.content for chunk in filtered_chunks])
        max_lexical = max(lexical_scores) if lexical_scores else 0.0

        ranked: list[tuple[float, float, float, int]] = []
        for idx, vector_score in enumerate(vector_scores):
            lexical_raw = lexical_scores[idx] if idx < len(lexical_scores) else 0.0
            lexical_norm = lexical_raw / max(max_lexical, 1e-9)
            final_score = (0.65 * float(vector_score)) + (0.35 * lexical_norm)
            ranked.append((final_score, float(vector_score), lexical_norm, idx))

        ranked.sort(key=lambda item: item[0], reverse=True)
        results: list[RAGDebugSearchResult] = []
        for final_score, vector_score, lexical_norm, idx in ranked[: max(k, 1)]:
            if min_score is not None and final_score < min_score:
                continue
            chunk = filtered_chunks[idx]
            preview = chunk.content.replace("\n", " ").strip()
            if len(preview) > 240:
                preview = f"{preview[:240]}..."
            metadata = dict(chunk.metadata) if chunk.metadata else {}
            metadata["vector_score"] = f"{vector_score:.6f}"
            metadata["lexical_score"] = f"{lexical_norm:.6f}"
            metadata["final_score"] = f"{final_score:.6f}"
            results.append(
                RAGDebugSearchResult(
                    chunk_id=chunk.chunk_id,
                    source=chunk.source,
                    score=final_score,
                    content_preview=preview,
                    metadata=metadata,
                )
            )
        return results


def _coerce_metadata(value: object) -> dict[str, str] | None:
    if not isinstance(value, dict):
        return None
    metadata: dict[str, str] = {}
    for key, item in value.items():
        if key is None:
            continue
        metadata[str(key)] = str(item)
    return metadata or None


def _match_source_filters(
    source: str,
    source_prefix: str | None = None,
    exclude_source_prefix: str | None = None,
) -> bool:
    if source_prefix and not source.startswith(source_prefix):
        return False
    if exclude_source_prefix and source.startswith(exclude_source_prefix):
        return False
    return True


def _is_memory_source(source: str) -> bool:
    return source.startswith(MEMORY_SOURCE_PREFIX)


def _slugify(value: str) -> str:
    normalized = re.sub(r"[^a-zA-Z0-9]+", "-", value.lower()).strip("-")
    return normalized or "chunk"


def _normalize_importance(value: str) -> str:
    lowered = value.strip().lower()
    if lowered in {"low", "medium", "high"}:
        return lowered
    return "medium"


def _read_document_text(path: Path, strict_quality: bool) -> tuple[str, list[str]]:
    extension = path.suffix.lower()
    if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_DOCUMENT_EXTENSIONS))
        raise RuntimeError(f"Unsupported file extension '{extension}'. Supported: {supported}")

    if extension in {".txt", ".md", ".markdown", ".rst"}:
        text = _read_plain_text_file(path)
    elif extension == ".pdf":
        text = _read_pdf_file(path)
    elif extension == ".docx":
        text = _read_docx_file(path)
    else:
        raise RuntimeError(f"Unsupported file extension '{extension}'.")

    warnings = _validate_extracted_text(path, text, strict_quality=strict_quality)
    return text, warnings


def _read_plain_text_file(path: Path) -> str:
    encodings = ["utf-8", "utf-8-sig"]
    for encoding in encodings:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise RuntimeError(
        f"Could not decode file as UTF-8 text: {path}. "
        "Binary files should use a supported parser extension (for example .pdf or .docx)."
    )


def _read_pdf_file(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Missing dependency 'pypdf'. Install project dependencies again.") from exc

    try:
        reader = PdfReader(str(path))
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(f"Could not open PDF file: {path}. Error: {exc}") from exc

    parts: list[str] = []
    for page in reader.pages:
        extracted = page.extract_text() or ""
        if extracted.strip():
            parts.append(extracted)
    return "\n\n".join(parts)


def _read_docx_file(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Missing dependency 'python-docx'. Install project dependencies again.") from exc

    try:
        doc = Document(str(path))
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(f"Could not open DOCX file: {path}. Error: {exc}") from exc

    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _validate_extracted_text(path: Path, text: str, strict_quality: bool) -> list[str]:
    normalized = text.strip()
    file_size = path.stat().st_size
    printable_ratio, replacement_chars = _text_quality(normalized)

    warnings: list[str] = []
    critical_issues: list[str] = []

    if not normalized:
        critical_issues.append("No text could be extracted from this file.")

    if replacement_chars > 0:
        warnings.append(f"Detected {replacement_chars} replacement characters (possible decode issue).")

    if printable_ratio < 0.80:
        critical_issues.append(
            f"Low printable ratio ({printable_ratio:.2f}). File may be binary/corrupted or badly extracted."
        )

    if file_size > 4096 and len(normalized) < 80:
        critical_issues.append(
            f"Very short extracted text ({len(normalized)} chars) for file size {file_size} bytes."
        )

    if strict_quality and critical_issues:
        message = " ".join(critical_issues)
        raise RuntimeError(f"Extraction quality check failed for {path}: {message}")

    warnings.extend(critical_issues)
    return warnings


def _text_quality(text: str) -> tuple[float, int]:
    if not text:
        return 1.0, 0
    non_space = [char for char in text if not char.isspace()]
    if not non_space:
        return 1.0, text.count("\ufffd")

    printable = sum(1 for char in non_space if char.isprintable())
    ratio = printable / max(len(non_space), 1)
    return ratio, text.count("\ufffd")


def _normalize_tokens(value: str) -> list[str]:
    cleaned = re.sub(r"[^a-zA-Z0-9À-ÿ]+", " ", value.lower())
    tokens = [token for token in cleaned.split() if token and token not in _STOPWORDS]
    return tokens


def _compute_lexical_scores(query: str, documents: list[str]) -> list[float]:
    query_tokens = _normalize_tokens(query)
    if not query_tokens:
        return [0.0 for _ in documents]

    query_counter = Counter(query_tokens)
    unique_query_terms = set(query_counter)
    scores: list[float] = []

    for document in documents:
        doc_tokens = _normalize_tokens(document)
        if not doc_tokens:
            scores.append(0.0)
            continue

        doc_counter = Counter(doc_tokens)
        matched_unique = len(unique_query_terms.intersection(doc_counter.keys()))
        coverage = matched_unique / max(len(unique_query_terms), 1)

        tf_sum = 0.0
        for term, q_count in query_counter.items():
            tf_sum += min(doc_counter.get(term, 0), 6) * q_count
        tf_norm = tf_sum / max(len(doc_tokens), 1)
        score = (0.70 * coverage) + (0.30 * tf_norm)
        scores.append(score)

    return scores
